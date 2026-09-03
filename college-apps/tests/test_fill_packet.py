#!/usr/bin/env python3
"""Unit tests for fill_packet.py (Post-Secondary Options Packet generator).

Tests verify that packet.json is converted into a valid Word document (.docx)
using templates/packet-template.docx, with correct table mappings, reflection
preservation, and [to be completed] placeholders for empty fields.
"""

import os
import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

try:
    from docx import Document
    import fill_packet as fp
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class FillPacketTest(unittest.TestCase):

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.out_path = Path(self.tmp.name) / "out" / "packet.docx"

    def tearDown(self):
        self.tmp.cleanup()

    def test_clean_build_with_full_data(self):
        data = {
            "name": "Jordan K",
            "email": "jordan@example.com",
            "phone": "555-123-4567",
            "high_school": "Roosevelt High School",
            "grad_year": 2027,
            "classes": [
                {"first": "AP Calculus BC", "second": "AP Calculus BC"},
                {"first": "AP Physics C", "second": "AP Physics C"}
            ],
            "teachers": [
                "Ms. Alvarez — AP Physics C, 11th",
                "Mr. Davis — AP English Language, 11th"
            ],
            "school_activities": [
                {"name": "Robotics Club", "grades": "10-12", "role": "Build Lead"}
            ],
            "outside_activities": [
                {"name": "Bike Clinic", "grades": "11-12", "role": "Volunteer Mechanic"}
            ],
            "hobbies": ["Tool building", "Cycling"],
            "honors": [
                {"name": "National Merit Commended", "grades": "11", "award": "Commended"}
            ],
            "work": [
                {"employer": "Garden Center", "grades": "11-12", "position": "Cashier"}
            ],
            "reflections": {
                "qualities": [
                    {"quality": "Intellectual Grit", "example": "Rotational motion retake"}
                ],
                "academic_growth": ["Learning how to study independently for physics"],
                "intellectual_growth": ["Reading beyond the curriculum in engineering"],
                "impact_campus": "Trained 6 freshmen on CNC machine",
                "impact_community": "Repaired 40+ bicycles for commuters",
                "challenges": "Worked 15 hrs/wk during junior year to support family",
                "challenges_include": "Yes",
                "majors": ["Mechanical Engineering", "Robotics"]
            },
            "colleges": [
                {
                    "name": "University of Michigan",
                    "decision_plan": "EA",
                    "deadline": "2026-11-01",
                    "counselor_letter": True,
                    "app_type": "Common App"
                }
            ],
            "parent_worksheet": {
                "1": "Jordan is intensely curious about how mechanical things work."
            }
        }

        blanks = fp.build(data, self.out_path)
        self.assertTrue(self.out_path.exists())

        doc = Document(str(self.out_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)

        self.assertIn("Jordan K", full_text)
        self.assertIn("Roosevelt High School", full_text)
        self.assertIn("Ms. Alvarez", full_text)
        self.assertIn("Rotational motion retake", full_text)
        self.assertIn("University of Michigan", table_text)
        self.assertIn("PARENT/GUARDIAN WORKSHEET", full_text)
        self.assertIn("Jordan is intensely curious", full_text)

    def test_missing_fields_render_to_be_completed(self):
        minimal_data = {
            "name": "Alex T",
            "grad_year": 2027
        }
        blanks = fp.build(minimal_data, self.out_path)
        self.assertGreater(blanks, 0)

        doc = Document(str(self.out_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn(fp.TODO, full_text)

    def test_challenges_include_flag_preserved(self):
        data = {
            "name": "Sam R",
            "reflections": {
                "challenges": "Family health issue sophomore year",
                "challenges_include": "No"
            }
        }
        fp.build(data, self.out_path)
        doc = Document(str(self.out_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Family health issue sophomore year", full_text)
        self.assertIn("No", full_text)


if __name__ == "__main__":
    unittest.main()

#!/usr/bin/env python3
"""Render a filled-in Post-Secondary Options Packet as .docx.

Mirrors the structure of the school's own packet so a counselor sees the format they
expect. Reads packet.json (produced by the counselor-package skill from profile.md) —
the model does the extraction, this script only does deterministic layout.

Usage:
    fill_packet.py students/maya-rodriguez
    fill_packet.py students/maya-rodriguez -o /tmp/packet.docx

packet.json shape — every key optional, missing ones render as "[to be completed]":

    {
      "name": "...", "email": "...", "phone": "...",
      "high_school": "...", "grad_year": 2027,
      "classes": [{"first": "AP Physics C", "second": "AP Physics C"}],
      "teachers": ["Ms. Alvarez — AP Physics, 11th", "..."],
      "school_activities":  [{"name": "...", "grades": "9-12", "role": "..."}],
      "outside_activities": [{"name": "...", "grades": "10-12", "role": "..."}],
      "hobbies": ["...", "..."],
      "honors":  [{"name": "...", "grades": "11", "award": "..."}],
      "work":    [{"employer": "...", "grades": "11-12", "position": "..."}],
      "reflections": {
        "qualities": [{"quality": "...", "example": "..."}],
        "academic_growth":     ["...", "..."],
        "intellectual_growth": ["...", "..."],
        "impact_campus": "...", "impact_community": "...",
        "challenges": "...", "challenges_include": "Yes" | "No",
        "majors": ["...", "..."]
      },
      "colleges": [ ... from meta.json ... ],
      "parent_worksheet": {"1": "...", ..., "7": "..."}
    }
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    sys.exit("python-docx not installed. Run: .venv/bin/pip install python-docx")

TODO = "[to be completed]"

PARENT_QUESTIONS = [
    "Describe your child in 3 or 4 adjectives and explain why you chose those words.",
    "What are some significant ways in which your child has changed/grown since entering high school?",
    "Describe a characteristic or accomplishment that we might not know about - something "
    "that won't show up on your child's activity record.",
    "What is something your child might not tell us, but you feel is important for us to know?",
    "What do you consider the most important achievement of your child during high school? "
    "Why did you select this as important?",
    "As a parent or guardian, please indicate any significant information that you feel "
    "would be helpful to include in your student's recommendation letter.",
    "How has your child overcome obstacles/adversity and demonstrated strength, courage, "
    "or resiliency?",
]


def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.underline, r.font.size = True, True, Pt(12)
    p.space_before, p.space_after = Pt(14), Pt(4)
    return p


def field(doc, label, value):
    p = doc.add_paragraph()
    r = p.add_run(f"{label} ")
    r.bold = True
    v = p.add_run(str(value) if value else TODO)
    if not value:
        v.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p


def table(doc, headers, rows, min_rows=1):
    """Bordered table; pads to min_rows so blanks stay visibly fillable."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True

    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val) if val else ""

    for _ in range(max(0, min_rows - len(rows))):
        cells = t.add_row().cells
        for i in range(len(headers)):
            r = cells[i].paragraphs[0].add_run(TODO)
            r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    return t


def q(doc, number, text, answer, italic_hint=None):
    p = doc.add_paragraph()
    r = p.add_run(f"{number}. {text}")
    r.bold = True
    if italic_hint:
        h = doc.add_paragraph()
        hr = h.add_run(italic_hint)
        hr.italic = True
        hr.font.size = Pt(9)
    a = doc.add_paragraph(style="List Bullet" if number else None)
    ar = a.add_run(answer if answer else TODO)
    if not answer:
        ar.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p


def build(d, out):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Post-Secondary Options Packet")
    tr.bold, tr.font.size = True, Pt(15)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(d.get("high_school", ""))
    sr.font.size = Pt(11)

    field(doc, "Name:", d.get("name"))
    field(doc, "Personal email:", d.get("email"))
    field(doc, "Cell phone:", d.get("phone"))

    heading(doc, "Senior Year Classes")
    table(doc, ["First semester classes", "Second semester classes (tentative)"],
          [[c.get("first"), c.get("second")] for c in d.get("classes", [])], min_rows=7)

    heading(doc, "Teachers who know you well")
    teachers = d.get("teachers", [])
    for i in range(max(3, len(teachers))):
        p = doc.add_paragraph(style="List Number")
        val = teachers[i] if i < len(teachers) else None
        r = p.add_run(val if val else TODO)
        if not val:
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    heading(doc, "School Activities (e.g., clubs, teams, sports, performing arts, etc.)")
    table(doc, ["Name of group", "Grade or age", "Responsibilities or office held"],
          [[a.get("name"), a.get("grades"), a.get("role")]
           for a in d.get("school_activities", [])], min_rows=5)

    heading(doc, "Outside Activities")
    doc.add_paragraph().add_run(
        "How do you spend your time outside of school? Examples: religious organization, "
        "community involvement, part-time job, babysitting, etc."
    ).italic = True
    table(doc, ["Name of group", "Grade or age", "Responsibilities (if applies)"],
          [[a.get("name"), a.get("grades"), a.get("role")]
           for a in d.get("outside_activities", [])], min_rows=4)

    heading(doc, "Hobbies and/or Favorite Activities")
    hobbies = d.get("hobbies", [])
    pairs = [hobbies[i:i + 2] for i in range(0, max(len(hobbies), 6), 2)]
    table(doc, ["", ""], [[p[0] if p else None, p[1] if len(p) > 1 else None]
                          for p in pairs], min_rows=3)

    heading(doc, "Honors and Awards in High School")
    table(doc, ["Name of Group or Organization", "Grade or Age", "Award Received"],
          [[h.get("name"), h.get("grades"), h.get("award")]
           for h in d.get("honors", [])], min_rows=5)

    heading(doc, "Work Experience")
    table(doc, ["Employer", "Grade or Age", "Position"],
          [[w.get("employer"), w.get("grades"), w.get("position")]
           for w in d.get("work", [])], min_rows=3)

    doc.add_page_break()
    heading(doc, "Reflections/Your Story")
    intro = doc.add_paragraph()
    intro.add_run(
        "We want to showcase what you stand out for — your strengths, accomplishments, "
        "and what makes you unique."
    ).italic = True

    refl = d.get("reflections", {})

    p = doc.add_paragraph()
    p.add_run("1. List three distinctive qualities that best describe you. For each "
              "quality, include a specific example.").bold = True
    quals = refl.get("qualities", [])
    for i in range(max(3, len(quals))):
        item = quals[i] if i < len(quals) else {}
        b = doc.add_paragraph(style="List Bullet")
        b.add_run(f"Quality {i + 1}: ").bold = True
        qr = b.add_run(item.get("quality") or TODO)
        if not item.get("quality"):
            qr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        e = doc.add_paragraph(style="List Bullet 2")
        e.add_run("Example: ").bold = True
        er = e.add_run(item.get("example") or TODO)
        if not item.get("example"):
            er.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.add_run("2. Describe how you have grown during your time in high school.").bold = True
    for label, key in [("Academic Growth", "academic_growth"),
                       ("Intellectual Growth", "intellectual_growth")]:
        b = doc.add_paragraph(style="List Bullet")
        b.add_run(f"{label}: ").bold = True
        examples = refl.get(key, [])
        for i in range(max(2, len(examples))):
            e = doc.add_paragraph(style="List Bullet 2")
            e.add_run("Example: ").bold = True
            val = examples[i] if i < len(examples) else None
            er = e.add_run(val or TODO)
            if not val:
                er.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Heading only — Q3's content lives entirely in the two sub-bullets below. Routing
    # it through q() with answer=None emitted a permanent "[to be completed]" above
    # answers that were in fact filled in, which teaches the student to ignore the
    # marker everywhere else in the packet.
    p = doc.add_paragraph()
    p.add_run("3. Describe specific ways you've had an impact, whether at school, in "
              "your community, or beyond.").bold = True
    for label, key in [("Impact on Campus", "impact_campus"),
                       ("Impact on Community", "impact_community")]:
        b = doc.add_paragraph(style="List Bullet")
        b.add_run(f"{label}: ").bold = True
        val = refl.get(key)
        r = b.add_run(val or TODO)
        if not val:
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    q(doc, 4, "Have you faced any personal challenges or outside circumstances that have "
              "affected your academic journey?", refl.get("challenges"))
    field(doc, "Include this in your letter of recommendation?",
          refl.get("challenges_include"))

    p = doc.add_paragraph()
    p.add_run("5. Intended College Major(s)/Career Goal(s)").bold = True
    majors = refl.get("majors", [])
    for i in range(max(3, len(majors))):
        b = doc.add_paragraph(style="List Bullet")
        val = majors[i] if i < len(majors) else None
        r = b.add_run(val or TODO)
        if not val:
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()
    heading(doc, "Colleges to which you are applying")
    table(doc,
          ["College Names", "Decision Type", "Due Date", "Counselor Letter?",
           "Application Type"],
          [[c.get("name"), c.get("decision_plan"), c.get("deadline"),
            "Yes" if c.get("counselor_letter") else "No", c.get("app_type")]
           for c in d.get("colleges", [])], min_rows=6)

    parent = d.get("parent_worksheet", {})
    if parent:
        doc.add_page_break()
        heading(doc, "PARENT/GUARDIAN WORKSHEET")
        for i, question in enumerate(PARENT_QUESTIONS, start=1):
            q(doc, i, question, parent.get(str(i)))

    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    # Count what actually rendered as a placeholder. The previous approach counted the
    # literal string "null" in the source JSON, which both missed absent keys and
    # double-counted nested nulls in the college list.
    blanks = sum(p.text.count(TODO) for p in doc.paragraphs)
    blanks += sum(c.text.count(TODO)
                  for t in doc.tables for row in t.rows for c in row.cells)
    return blanks


def main():
    p = argparse.ArgumentParser(description=__doc__,
                                formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("student_dir")
    p.add_argument("-o", "--out")
    args = p.parse_args()

    sd = Path(args.student_dir)
    pj = sd / "packet.json"
    if not pj.exists():
        sys.exit(f"No packet.json in {sd}.\n"
                 "The counselor-package skill builds it from profile.md first.")
    d = json.loads(pj.read_text())

    meta = sd / "meta.json"
    if meta.exists() and "colleges" not in d:
        d["colleges"] = json.loads(meta.read_text()).get("colleges", [])

    out = args.out or sd / "out" / "packet.docx"
    blanks = build(d, out)

    print(f"Wrote {out}")
    if blanks:
        print(f"  note: {blanks} field(s) still empty — they render as '{TODO}'")


if __name__ == "__main__":
    main()
